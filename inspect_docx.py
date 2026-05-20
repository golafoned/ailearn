import os
from docx import Document
from docx.shared import Cm

def inspect():
    path = 'processed.docx'
    if not os.path.exists(path):
        print(f"File {path} does not exist.")
        return

    size = os.path.getsize(path)
    print(f"Existence: True, Size: {size} bytes")

    try:
        doc = Document(path)
    except Exception as e:
        print(f"Error opening document: {e}")
        return

    # 2) Paragraphs with '$'
    dollars = [p.text for p in doc.paragraphs if '$' in p.text]
    print(f"Paragraphs with '$': {len(dollars)}")

    # 3) Match equations
    eq1_target = 'M_{new} = \\min(100, \\max(0, M_{old} + w(d, r)))'
    eq2_target = "EF' = EF + \\left(0.1 - (5 - q)\\left(0.08 + (5 - q) \\cdot 0.02\\right)\\right)"
    
    found_eq1 = [p.text for p in doc.paragraphs if eq1_target in p.text]
    found_eq2 = [p.text for p in doc.paragraphs if eq2_target in p.text]
    
    print(f"Equation 1 matches: {found_eq1}")
    print(f"Equation 2 matches: {found_eq2}")

    # 4) Images
    image_count = 0
    max_width_cm = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1
    
    # Inline shapes for width
    for shape in doc.inline_shapes:
        try:
            width_cm = shape.width.cm
            if width_cm > max_width_cm:
                max_width_cm = width_cm
        except:
            pass
            
    print(f"Image count: {image_count}")
    print(f"Max image width: {max_width_cm:.2f} cm")

if __name__ == "__main__":
    inspect()
