import os
import re
import tempfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
import matplotlib
from matplotlib import mathtext

matplotlib.use('Agg')

FORMULA_SPECS = [
    {
        'source': r"M_{new} = \min(100, \max(0, M_{old} + w(d, r)))",
        'width_cm': 10.5,
    },
    {
        'source': r"EF' = EF + \left(0.1 - (5 - q)\left(0.08 + (5 - q) \cdot 0.02\right)\right)",
        'width_cm': 12.0,
    },
]

def normalize_formula(formula):
    return (
        formula.replace(r'\left', '')
        .replace(r'\right', '')
        .replace(r'\min', r'\mathrm{min}')
        .replace(r'\max', r'\mathrm{max}')
    )


def render_equation_image(formula, image_path):
    mathtext.math_to_image(f"${normalize_formula(formula)}$", image_path, dpi=300, format='png')


def clear_paragraph(paragraph):
    paragraph_element = paragraph._element
    for child in list(paragraph_element):
        if child.tag.endswith('pPr'):
            continue
        paragraph_element.remove(child)


def patch_document():
    input_path = 'processed.docx'
    output_path = 'coursework.docx'
    if os.path.exists(output_path):
        os.remove(output_path)

    doc = Document(input_path)
    temp_paths = []
    replaced = 0
    try:
        for spec in FORMULA_SPECS:
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
            temp_file.close()
            temp_paths.append(temp_file.name)
            render_equation_image(spec['source'], temp_file.name)

            for paragraph in doc.paragraphs:
                if paragraph.text.strip() != spec['source']:
                    continue
                clear_paragraph(paragraph)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(temp_file.name, width=Cm(spec['width_cm']))
                replaced += 1
                break

        doc.save(output_path)
    finally:
        for path in temp_paths:
            if os.path.exists(path):
                os.remove(path)
    return replaced


def validate():
    output_path = 'coursework.docx'
    doc = Document(output_path)
    paragraph_has_dollar = any('$' in para.text for para in doc.paragraphs)
    raw_formula_remaining = sum(
        1 for para in doc.paragraphs for spec in FORMULA_SPECS if para.text.strip() == spec['source']
    )
    image_count = len(doc.inline_shapes)
    max_width_cm = 0.0
    for shape in doc.inline_shapes:
        max_width_cm = max(max_width_cm, shape.width.cm)
    return {
        'exists': os.path.exists(output_path),
        'size': os.path.getsize(output_path),
        'paragraphHasDollar': paragraph_has_dollar,
        'rawFormulaRemaining': raw_formula_remaining,
        'imageCount': image_count,
        'maxWidthCm': round(max_width_cm, 2),
    }


if __name__ == '__main__':
    replaced = patch_document()
    results = validate()
    print(f'replaced={replaced}')
    for key, value in results.items():
        print(f'{key}={value}')
