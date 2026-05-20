import os
import zipfile
from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement, ns
from docx.shared import Cm

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def insert_toc(paragraph):
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(ns.qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(ns.qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(ns.qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(ns.qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

doc = Document("coursework.docx")

# 2) Different first page
for section in doc.sections:
    section.different_first_page_header_footer = True

# 1) Page numbers in top-right header
# We need to add it to the header (default header, which is for non-first pages if different_first_page is true)
section = doc.sections[0]
header = section.header
header.is_linked_to_previous = False
# Clear header or check if exists
header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
header_para.alignment = 1 # Center or 2 for right? User said top-right. 2 is RIGHT.
from docx.enum.text import WD_ALIGN_PARAGRAPH
header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_page_number(header_para.add_run())

# 3) TOC field under 'ЗМІСТ'
target_text = 'ЗМІСТ'
for i, para in enumerate(doc.paragraphs):
    if target_text in para.text:
        # Insert TOC after this paragraph
        p = doc.paragraphs[i]
        new_p = doc.add_paragraph() # Temporary
        # Move new_p after p
        p._p.addnext(new_p._p)
        insert_toc(new_p)
        break

doc.save("coursework.docx")

# Validation
doc = Document("coursework.docx")
file_size = os.path.getsize("coursework.docx")
para_count = len(doc.paragraphs)
inline_images = doc.inline_shapes
img_count = len(inline_images)
max_width_cm = 0
for img in inline_images:
    w = img.width.cm if img.width else 0
    if w > max_width_cm:
        max_width_cm = w

has_dollar = any("$" in p.text for p in doc.paragraphs)
# Raw display formula paragraphs (OMath)
has_omath = False
for p in doc.paragraphs:
    if p._element.xpath(".//m:oMath"):
        has_omath = True
        break

# XML checks
has_toc_xml = False
has_page_xml = False
diff_first_xml = False

with zipfile.ZipFile("coursework.docx", "r") as z:
    doc_xml = z.read("word/document.xml").decode("utf-8")
    if 'TOC \\o "1-3"' in doc_xml:
        has_toc_xml = True
    
    # Check headers
    for name in z.namelist():
        if name.startswith("word/header"):
            h_xml = z.read(name).decode("utf-8")
            if "PAGE" in h_xml:
                has_page_xml = True
    
    # Check setting
    if 'w:titlePg' in doc_xml:
        diff_first_xml = True

print(f"File Size: {file_size}")
print(f"Paragraph Count: {para_count}")
print(f"Inline Image Count: {img_count}")
print(f"Max Image Width (cm): {max_width_cm:.2f}")
print(f"Paragraphs with '$': {has_dollar}")
print(f"OMath present: {has_omath}")
print(f"TOC field in XML: {has_toc_xml}")
print(f"PAGE field in Header XML: {has_page_xml}")
print(f"Different-First-Page XML: {diff_first_xml}")
